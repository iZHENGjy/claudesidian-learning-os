#!/usr/bin/env python3
"""process_si.py — SI 多格式 dispatcher

输入 (YYYY) Title 文件夹，自动找 si.<ext> 处理：
  - .pdf   → 调 mineru_convert（同 main.pdf 流程，输出 si.md + images_si/）
  - .docx  → python-docx 转 md（段落 + 表格 + 内嵌图）
  - .xlsx  → pandas 把每个 sheet 转 md table
  - .zip   → 解压后递归处理（找里面的 .pdf/.docx/.xlsx）
  - 其他   → 原文件存档进 _attachments_orig/，si.md 加占位

用法：
    py .scripts/process_si.py "<paper-folder>"
    py .scripts/process_si.py "<paper-folder>" --force
"""
from __future__ import annotations
import argparse
import io
import re
import sys
import shutil
import zipfile
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")

sys.path.insert(0, str(Path(__file__).parent))
from rich.console import Console

console = Console()
ARCHIVE_DIR = "_attachments_orig"


# ============ 入口分发 ============
def find_si_files(paper_dir: Path) -> list[Path]:
    """找 paper_dir 下所有 si.* 文件（排除已转好的 si.md）。"""
    out = []
    for p in paper_dir.glob("si.*"):
        if p.suffix.lower() == ".md":
            continue
        out.append(p)
    return out


# ============ PDF SI（调 mineru_convert）============
def process_pdf(si_pdf: Path, paper_dir: Path, force: bool = False) -> None:
    from mineru_convert import convert_via_api
    paper_id = re.sub(r"[^\w\-]", "_", paper_dir.name)[:80]
    convert_via_api(si_pdf, paper_dir, "si", paper_id, force=force)


# ============ DOCX SI ============
def process_docx(si_docx: Path, paper_dir: Path, force: bool = False) -> None:
    """python-docx 转 md。段落顺序保留 + 表格统一附在末尾（简化）+ 内嵌图抽到 images_si/。"""
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("缺包：pip install python-docx")

    target_md = paper_dir / "si.md"
    if target_md.exists() and not force:
        console.print(f"[dim]si.md 已存在，跳过（--force 覆盖）[/]")
        return

    images_dir = paper_dir / "images_si"
    images_dir.mkdir(exist_ok=True)

    doc = Document(si_docx)
    md = ["# Supporting Information", ""]

    # 段落（顺序保留，按 style 转 heading）
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        style = (p.style.name or "").lower() if p.style else ""
        if "heading 1" in style:
            md.append(f"## {text}")
        elif "heading 2" in style:
            md.append(f"### {text}")
        elif "heading 3" in style or "heading 4" in style:
            md.append(f"#### {text}")
        else:
            md.append(text)
        md.append("")

    # 表格（追加在末尾，简化版）
    if doc.tables:
        md.append("\n## 表格\n")
        for i, tbl in enumerate(doc.tables, 1):
            md.append(f"### Table {i}")
            md.append(_docx_table_to_md(tbl))
            md.append("")

    # 内嵌图
    img_idx = 0
    for rel in doc.part.rels.values():
        if "image" in (rel.target_ref or ""):
            try:
                blob = rel.target_part.blob
            except Exception:
                continue
            img_idx += 1
            ext = Path(rel.target_ref).suffix or ".png"
            out_name = f"figS{img_idx}_inline{ext}"
            (images_dir / out_name).write_bytes(blob)

    if img_idx > 0:
        md.append(f"\n## 内嵌图 ({img_idx} 张)\n")
        for i in range(1, img_idx + 1):
            ext = ".png"  # 大多数 docx 嵌入图是 png
            md.append(f"![](images_si/figS{i}_inline{ext})")
            md.append("")

    target_md.write_text("\n".join(md), encoding="utf-8")
    console.print(f"[green]✓ docx → {target_md.name}[/] "
                  f"({len(doc.paragraphs)} 段 / {len(doc.tables)} 表 / {img_idx} 内嵌图)")


def _docx_table_to_md(tbl) -> str:
    """word 表格 → markdown table（简化版，不处理合并单元格）。"""
    rows = []
    for row in tbl.rows:
        cells = []
        for cell in row.cells:
            text = cell.text.replace("|", "\\|").replace("\n", " ").strip()
            cells.append(text)
        rows.append("| " + " | ".join(cells) + " |")
    if not rows:
        return "(空表)"
    # 在第一行后插入分隔
    n_cols = len(tbl.rows[0].cells) if tbl.rows else 0
    sep = "| " + " | ".join(["---"] * n_cols) + " |"
    rows.insert(1, sep)
    return "\n".join(rows)


# ============ XLSX SI ============
def process_xlsx(si_xlsx: Path, paper_dir: Path, force: bool = False) -> None:
    try:
        import pandas as pd
    except ImportError:
        raise RuntimeError("缺包：pip install pandas openpyxl")

    target_md = paper_dir / "si.md"
    if target_md.exists() and not force:
        console.print(f"[dim]si.md 已存在，跳过[/]")
        return

    md = ["# Supporting Information (xlsx)", ""]
    sheets = pd.read_excel(si_xlsx, sheet_name=None, engine="openpyxl")
    for sheet_name, df in sheets.items():
        md.append(f"## Sheet: {sheet_name}")
        if df.empty:
            md.append("(空 sheet)")
        else:
            md.append(df.to_markdown(index=False))
        md.append("")

    target_md.write_text("\n".join(md), encoding="utf-8")
    console.print(f"[green]✓ xlsx → {target_md.name}[/] ({len(sheets)} 个 sheet)")


# ============ ZIP SI ============
def process_zip(si_zip: Path, paper_dir: Path, force: bool = False) -> None:
    """解压找 .pdf/.docx/.xlsx 递归处理。多份则合并到一份 si.md。"""
    tmp = paper_dir / "_tmp_si_unzip"
    if tmp.exists():
        shutil.rmtree(tmp)
    tmp.mkdir()
    with zipfile.ZipFile(si_zip) as zf:
        zf.extractall(tmp)

    pdfs = list(tmp.rglob("*.pdf"))
    docxs = list(tmp.rglob("*.docx"))
    xlsxs = list(tmp.rglob("*.xlsx"))
    found = pdfs + docxs + xlsxs

    if not found:
        console.print(f"[yellow]⚠️ zip 里没找到 .pdf / .docx / .xlsx[/]")
        archive_orig(si_zip, paper_dir, "zip 内容是不可识别格式")
        shutil.rmtree(tmp)
        return

    # 简化：找一个最大的 PDF 当 si.pdf 跑（最常见情况）
    if pdfs:
        biggest_pdf = max(pdfs, key=lambda p: p.stat().st_size)
        console.print(f"[dim]zip 里找到 {len(pdfs)} 个 PDF，取最大的当 si.pdf：{biggest_pdf.name}[/]")
        target_si_pdf = paper_dir / "si.pdf"
        if target_si_pdf.exists() and not force:
            console.print(f"[yellow]si.pdf 已存在，跳过[/]")
        else:
            shutil.copy(biggest_pdf, target_si_pdf)
            process_pdf(target_si_pdf, paper_dir, force=force)
    elif docxs:
        process_docx(docxs[0], paper_dir, force=force)
    elif xlsxs:
        process_xlsx(xlsxs[0], paper_dir, force=force)

    shutil.rmtree(tmp)


# ============ 兜底：存档原文件 ============
def archive_orig(orig: Path, paper_dir: Path, reason: str = "") -> None:
    archive_dir = paper_dir / ARCHIVE_DIR
    archive_dir.mkdir(exist_ok=True)
    target = archive_dir / orig.name
    if not target.exists():
        shutil.move(str(orig), str(target))
    si_md = paper_dir / "si.md"
    placeholder = (
        f"# Supporting Information\n\n"
        f"> SI 是非可转换格式（{orig.suffix or '无后缀'}），原文件存档于 "
        f"[{orig.name}]({ARCHIVE_DIR}/{orig.name})\n"
    )
    if reason:
        placeholder += f"> 原因：{reason}\n"
    if not si_md.exists():
        si_md.write_text(placeholder, encoding="utf-8")
    console.print(f"[yellow]⚠️ 存档：{orig.name} → {ARCHIVE_DIR}/[/]")


# ============ main ============
def main():
    parser = argparse.ArgumentParser(description=__doc__.split("\n")[0],
                                     formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument("paper_dir", help="(YYYY) Title 文件夹路径")
    parser.add_argument("--force", action="store_true", help="覆盖已有 si.md")
    args = parser.parse_args()

    paper_dir = Path(args.paper_dir).resolve()
    if not paper_dir.is_dir():
        console.print(f"[red]❌ 不是目录: {paper_dir}[/]")
        sys.exit(1)

    si_files = find_si_files(paper_dir)
    if not si_files:
        console.print(f"[yellow]⚠️ {paper_dir.name} 里没找到 si.* 文件[/]")
        return

    console.print(f"\n[bold cyan]处理 {paper_dir.name} 的 SI（{len(si_files)} 个文件）[/]")

    handlers = {
        ".pdf": process_pdf,
        ".docx": process_docx,
        ".xlsx": process_xlsx,
        ".zip": process_zip,
    }

    for si in si_files:
        ext = si.suffix.lower()
        console.print(f"\n→ {si.name}（{ext}）")
        handler = handlers.get(ext)
        try:
            if handler:
                handler(si, paper_dir, force=args.force)
            else:
                archive_orig(si, paper_dir, f"{ext} 不在支持列表内")
        except Exception as e:
            console.print(f"[red]❌ 处理 {si.name} 失败: {e}[/]")


if __name__ == "__main__":
    main()
