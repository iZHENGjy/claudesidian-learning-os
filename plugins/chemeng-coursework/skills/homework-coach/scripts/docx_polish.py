"""homework-coach docx 后处理工具。

跑这个之前先让 pandoc 把 draft.md 转成 docx, 然后用本脚本兜底中文字体 / 表格 / 图片居中。

用法:
    # 1. pandoc 转 docx (公式必须加 tex_math_dollars)
    pandoc draft.md -o final.docx --from markdown+tex_math_dollars

    # 2. polish (默认三线表)
    python docx_polish.py final.docx

    # 想要旧的全网格表格 (Table Grid) 而不是三线表:
    python docx_polish.py final.docx --grid

做的事:
- 给所有 run 设 East Asia 字体为 Microsoft YaHei (兜底中文乱字体), 英文/数字保留 Times New Roman
- 表格 (默认): **三线表** (顶线 + 表头下线 + 底线, 无竖线无内横线) + **固定 DXA 内容自适应列宽** (满页宽,
  不挤在左边窄栏; 不用百分比——百分比在 Google Docs 会崩) + 9pt 表格字 + 表头加粗居中 + 单元格垂直居中
- 含图片的段落居中

不做:
- 改 margin / 段距 / 正文字号 (这些由 pandoc 默认 reference docx 控制, 如要改用学校模板 --reference-doc=...)
- 改公式 (pandoc tex_math_dollars 已经转成 OMML, 不动)

注: 列宽按 A4 (210mm) 减 2.54cm 边距 = 9026 twips 算。换 Letter/不同边距时改 USABLE_TWIPS。
"""

from __future__ import annotations
import io
import math
import sys
from pathlib import Path

# Windows 控制台默认 GBK, 强制 utf-8 防中文/特殊字符报错
try:
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
except Exception:
    pass

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("缺 python-docx, 跑: pip install python-docx", file=sys.stderr)
    sys.exit(1)


CHINESE_FONT = "Microsoft YaHei"
WESTERN_FONT = "Times New Roman"

# 三线表 / 列宽参数
USABLE_TWIPS = 9026   # A4(11906) - 2×2.54cm 边距(2880); 1 inch = 1440 twips
MIN_COL = 620         # 列最小宽 ~0.43", 防某列被压成 0
CAP_FRAC = 0.30       # 单列封顶 = 页宽的 30%, 防长表头独占
TABLE_PT = 9          # 表格字号


def set_run_fonts(run) -> None:
    """给 run 设字体: 中文走 East Asia, 英文/数字走 Western。

    python-docx 默认只设 Western, 中文会 fallback 到 docx 默认 (常常变宋体或方框)。
    必须手动给 rFonts 元素加 w:eastAsia 属性。
    """
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), CHINESE_FONT)
    rFonts.set(qn("w:ascii"), WESTERN_FONT)
    rFonts.set(qn("w:hAnsi"), WESTERN_FONT)


def polish_paragraphs(doc) -> int:
    """正文段落字体兜底。返回处理的 run 数。"""
    n = 0
    for para in doc.paragraphs:
        for run in para.runs:
            set_run_fonts(run)
            n += 1
    return n


def _set(parent, tag, **attrs):
    """get_or_add 一个子元素并设属性 (属性名自动加 w: 前缀)。"""
    el = parent.find(qn(tag))
    if el is None:
        el = OxmlElement(tag)
        parent.append(el)
    for k, v in attrs.items():
        el.set(qn("w:" + k), str(v))
    return el


def _col_widths(table):
    """内容自适应固定列宽 (twips), 和 = 页宽。

    权重以**数据格**最长文本为主, 表头按 0.45 折算 (表头能换行, 不该独占宽度),
    再对单列封顶, 避免某一列吃掉太多。
    """
    rows = table.rows
    if not rows:
        return []
    ncol = len(table.columns)
    weights = []
    for j in range(ncol):
        head = len(rows[0].cells[j].text.strip())
        data = 2
        for row in rows[1:]:
            try:
                data = max(data, len(row.cells[j].text.strip()))
            except IndexError:
                pass
        weights.append(max(data, math.ceil(head * 0.45), 2))
    total = sum(weights) or 1
    raw = [max(MIN_COL, int(USABLE_TWIPS * w / total)) for w in weights]
    cap = int(USABLE_TWIPS * CAP_FRAC)
    raw = [min(cap, x) for x in raw]
    s = sum(raw) or 1
    widths = [int(x * USABLE_TWIPS / s) for x in raw]
    widths[-1] += USABLE_TWIPS - sum(widths)   # 误差补最后一列
    return widths


def _three_line_borders(table) -> None:
    """三线表: 顶线 + 表头下线 + 底线, 无竖线无内横线 (booktabs 学术风)。"""
    tblPr = table._tbl.tblPr
    old = tblPr.find(qn("w:tblBorders"))
    if old is not None:
        tblPr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom"):       # 顶/底线 1.5pt (sz=12)
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "12")
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), "000000")
        borders.append(e)
    for edge in ("left", "right", "insideH", "insideV"):  # 去竖线 + 内横线
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), "nil")
        borders.append(e)
    tblPr.append(borders)
    # 表头行下线 0.75pt (sz=6)
    if table.rows:
        for cell in table.rows[0].cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcB = tcPr.find(qn("w:tcBorders"))
            if tcB is None:
                tcB = OxmlElement("w:tcBorders"); tcPr.append(tcB)
            ob = tcB.find(qn("w:bottom"))
            if ob is not None:
                tcB.remove(ob)
            b = OxmlElement("w:bottom")
            b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "6")
            b.set(qn("w:space"), "0"); b.set(qn("w:color"), "000000")
            tcB.append(b)


def _style_table_runs(table) -> None:
    """表头加粗居中 + 全表字体兜底 + 9pt。"""
    if not table.rows:
        return
    for cell in table.rows[0].cells:        # 表头
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(TABLE_PT)
                set_run_fonts(run)
    for row in table.rows[1:]:              # 其余
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(TABLE_PT)
                    set_run_fonts(run)


def polish_tables(doc, three_line: bool = True) -> int:
    """表格美化。three_line=True (默认): 三线表 + 固定 DXA 满页宽列宽 + 9pt。
    three_line=False: 旧的 Table Grid (全网格)。返回表数。"""
    n = 0
    for table in doc.tables:
        if len(table.rows) == 0:
            continue
        if three_line:
            table.autofit = False
            table.allow_autofit = False
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            _three_line_borders(table)
            tblPr = table._tbl.tblPr
            _set(tblPr, "w:tblLayout", type="fixed")
            _set(tblPr, "w:tblW", type="dxa", w=USABLE_TWIPS)
            widths = _col_widths(table)
            grid = table._tbl.find(qn("w:tblGrid"))
            if grid is not None:
                for j, gc in enumerate(grid.findall(qn("w:gridCol"))):
                    if j < len(widths):
                        gc.set(qn("w:w"), str(widths[j]))
                        gc.set(qn("w:type"), "dxa")
            for row in table.rows:          # fixed 布局下每格也要固定宽
                for j, cell in enumerate(row.cells):
                    if j < len(widths):
                        _set(cell._tc.get_or_add_tcPr(), "w:tcW",
                             type="dxa", w=widths[j])
            _style_table_runs(table)
        else:
            try:
                table.style = "Table Grid"
            except KeyError:
                pass
            _style_table_runs(table)
        n += 1
    return n


def center_figures(doc) -> int:
    """所有包含图片 (w:drawing) 的段落居中。返回处理的段数。"""
    n = 0
    for para in doc.paragraphs:
        has_image = any(
            run.element.findall(qn("w:drawing"))
            or run.element.findall(qn("w:pict"))
            for run in para.runs
        )
        if has_image:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            n += 1
    return n


def main() -> int:
    args = [a for a in sys.argv[1:] if not a.startswith("-")]
    three_line = "--grid" not in sys.argv
    if len(args) != 1:
        print(f"用法: python {sys.argv[0]} <docx 路径> [--grid]", file=sys.stderr)
        return 2

    docx_path = Path(args[0])
    if not docx_path.exists():
        print(f"文件不存在: {docx_path}", file=sys.stderr)
        return 1

    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"打不开 docx (是不是被 Word 占用了?): {e}", file=sys.stderr)
        return 1

    n_runs = polish_paragraphs(doc)
    n_tables = polish_tables(doc, three_line=three_line)
    n_figs = center_figures(doc)

    try:
        doc.save(docx_path)
    except PermissionError:
        print(f"写不进 {docx_path}, 多半是 Word/阅读器打开了, 关掉再 retry", file=sys.stderr)
        return 1

    style = "三线表 DXA" if three_line else "Table Grid"
    print(
        f"[OK] polished {docx_path.name}: "
        f"{n_runs} runs 字体兜底, {n_tables} 张表({style}), {n_figs} 张图居中"
    )
    return 0


if __name__ == "__main__":
    sys.exit(main())
