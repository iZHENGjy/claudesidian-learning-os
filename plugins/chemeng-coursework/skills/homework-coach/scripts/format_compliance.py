# -*- coding: utf-8 -*-
"""课程报告提交格式合规后处理 (在 docx_polish.py 之后跑, 最后一步)。
这是 chemeng-coursework 插件的**共享版本**——每份课程报告 build 都该指向它,
踩到新格式坑就往这里加规则, 别只改单份作业 (治"格式默认值每次重新发现"病)。

治课程说明常见硬性格式要求:
  1. 正文 Times New Roman 10pt / 两端对齐 / 单倍行距 / 段后 6pt / 零缩进
  2. 显示公式 → 居中 + 右对齐编号 (1)(2)...
  3. 三栏页眉: 姓名(左) | 学号(中) | 课程代码(右)
  4. 页脚: 居中粗体页码
  5. 代码块 → Courier New 8pt (等宽; 9pt 会让长行断词折行)
  6. 化学计量表列宽重分配 (Outlet 列加宽防换行)
  7. 标题一律黑色 ← pandoc 默认用 Word 内置 Heading 样式=蓝色主题色, 学术报告要黑

⚠️ 这些数值 (10pt / TNR / 段后 6pt / 页眉三栏) 是照一份课程说明定死的,
   不是通用标准。你的课要求不一样就直接改下面的常量。

跑法: python format_compliance.py final.docx [COURSE_CODE]
     COURSE_CODE 不给就在页眉留占位符, 跟 [Name]/[Student ID] 一样自己填。
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
import docx
from docx.shared import Pt, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BODY_STYLES = {"Body Text", "First Paragraph", "Compact"}   # 正文段落样式
HEADING_STYLES = {"Title", "Heading 1", "Heading 2", "Heading 3",
                  "Heading 4", "Heading 5"}                 # 要强制黑色的标题
TNR = "Times New Roman"
BLACK = RGBColor(0, 0, 0)
COURSE = "[Course Code]"   # 页眉右栏课程代码; 命令行第 2 个参数可覆盖

def set_run(r, name=TNR, size=None, bold=None, color=None):
    """强制设 run 的西文+东亚字体, 可选字号/加粗/颜色。"""
    r.font.name = name
    rpr = r._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(a), name)
    if size is not None:
        r.font.size = Pt(size)
    if bold is not None:
        r.font.bold = bold
    if color is not None:
        r.font.color.rgb = color

def blacken_heading_styles(d):
    """样式层把标题颜色设成黑 (盖掉 pandoc 继承的蓝色 themeColor)。"""
    for sn in HEADING_STYLES:
        try:
            d.styles[sn].font.color.rgb = BLACK
        except KeyError:
            pass

def usable_width_twips(section):
    """页面可用宽度 (twips) = 页宽 - 左右边距。pandoc 常不写尺寸→None, 用 Letter 1in 兜底。"""
    pw = section.page_width or Twips(12240)      # Letter 8.5in
    lm = section.left_margin or Twips(1440)      # 1in
    rm = section.right_margin or Twips(1440)
    return int((pw - lm - rm) / 635)             # EMU->twips

def add_page_field(paragraph):
    """给段落加一个 PAGE 域 (自动页码)。"""
    run = paragraph.add_run()
    r = run._element
    for t, attr in (('begin', None), (None, 'PAGE'), ('end', None)):
        if t:
            fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), t); r.append(fc)
        else:
            it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = ' PAGE '; r.append(it)
    return run

def main(path, course=COURSE):
    d = docx.Document(path)
    sec = d.sections[0]
    W = usable_width_twips(sec)

    # ---- 0: 标题颜色统一黑 (样式层, 盖掉蓝色 themeColor) ----
    blacken_heading_styles(d)

    # ---- 1 + 2 + 5 + 7: 逐段处理 ----
    eq_n = 0
    body_cnt = 0
    head_cnt = 0
    for p in d.paragraphs:
        style = p.style.name or ""
        has_math = bool(p._p.findall('.//' + qn('m:oMath')))
        is_display_eq = has_math and p.text.strip() == ""

        if is_display_eq:
            # 显示公式: tab stop 让公式居中 + 编号右对齐
            eq_n += 1
            pf = p.paragraph_format
            pf.tab_stops.add_tab_stop(Twips(W // 2), WD_TAB_ALIGNMENT.CENTER)  # 公式居中
            pf.tab_stops.add_tab_stop(Twips(W), WD_TAB_ALIGNMENT.RIGHT)        # 编号靠右
            ppr = p._p.find(qn('w:pPr'))
            tr = OxmlElement('w:r'); tr.append(OxmlElement('w:tab'))
            if ppr is not None:
                ppr.addnext(tr)
            else:
                p._p.insert(0, tr)
            run = p.add_run('\t(%d)' % eq_n)
            set_run(run, TNR, 10)
            continue

        if style in HEADING_STYLES:
            # 标题: 每个 run 也强制黑 + TNR (run 层覆盖, 双保险), 保留原字号/粗体
            for r in p.runs:
                set_run(r, TNR, color=BLACK)
            head_cnt += 1
            continue

        if style == "Source Code":
            # 代码: 等宽 8pt (9pt 长行会断词折行, 8pt 完整不折行更清晰)
            for r in p.runs:
                set_run(r, "Courier New", 8)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            continue

        if style in BODY_STYLES:
            # 正文: TNR 10 / 两端对齐 / 单倍 / 段后6 / 零缩进
            for r in p.runs:
                set_run(r, TNR, 10)
            pf = p.paragraph_format
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            pf.space_after = Pt(6)
            pf.first_line_indent = Pt(0)
            pf.left_indent = Pt(0)
            body_cnt += 1

    print(f"[OK] 正文 {body_cnt} 段 TNR10; 标题 {head_cnt} 段统一黑色; 显示公式编号 1-{eq_n}")

    # ---- 3: 三栏页眉 (姓名 | 学号 | 课程代码) ----
    hp = sec.header.paragraphs[0]
    hp.text = ""
    hp.paragraph_format.tab_stops.add_tab_stop(Twips(W // 2), WD_TAB_ALIGNMENT.CENTER)
    hp.paragraph_format.tab_stops.add_tab_stop(Twips(W), WD_TAB_ALIGNMENT.RIGHT)
    hr = hp.add_run("[Name]\t[Student ID]\t%s" % course)
    set_run(hr, TNR, 9)
    print(f"[OK] 三栏页眉: [Name] | [Student ID] | {course} (占位符待填)")

    # ---- 4: 页脚居中粗体页码 ----
    fp = sec.footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr = add_page_field(fp)
    set_run(pr, TNR, 10, bold=True)
    print("[OK] 页脚: 居中粗体自动页码")

    # ---- 6: 化学计量表列宽重分配 (Concentration 分数其实窄, 匀给 Outlet 防换行) ----
    # 只动表头首格 = "Species" 的表 (化学计量表), 不碰数据表/评分表/声明表
    W5 = [1368, 1224, 1440, 2688, 2640]   # 5列: Species|Feed|Change|Outlet|Conc
    W4 = [2016, 1584, 1584, 4176]         # 4列: Species|Feed|Change|Conc
    fixed = 0
    for tbl in d.tables:
        cells0 = tbl.rows[0].cells
        if not cells0 or cells0[0].text.strip() != "Species":
            continue
        widths = W5 if len(tbl.columns) == 5 else (W4 if len(tbl.columns) == 4 else None)
        if widths is None:
            continue
        tbl.autofit = False
        for row in tbl.rows:
            for j, cell in enumerate(row.cells):
                if j < len(widths):
                    cell.width = Twips(widths[j])
        fixed += 1
    print(f"[OK] 化学计量表列宽重分配: {fixed} 张 (Outlet 列加宽防换行)")

    d.save(path)
    print(f"[DONE] 已写回 {path}")

if __name__ == "__main__":
    path = sys.argv[1] if len(sys.argv) > 1 else "final.docx"
    course = sys.argv[2] if len(sys.argv) > 2 else COURSE
    main(path, course)
